"""
File parsing service for extracting text from PDF and DOCX files.

Supports:
    - PDF via pypdf
    - DOCX via python-docx
"""

import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract plain text from a PDF file.

    Parameters
    ----------
    file_bytes : bytes
        Raw bytes of the PDF file.

    Returns
    -------
    str
        Extracted text content.

    Raises
    ------
    ValueError
        If the PDF cannot be read or is encrypted.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))

        if reader.is_encrypted:
            raise ValueError("Cannot extract text from an encrypted PDF. Please provide an unencrypted file.")

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())

        extracted = "\n\n".join(text_parts)

        if not extracted.strip():
            raise ValueError("No text could be extracted from this PDF. It may be a scanned image-only document.")

        logger.info(f"Extracted {len(extracted)} characters from PDF ({len(reader.pages)} pages).")
        return extracted

    except ImportError:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}") from e


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract plain text from a DOCX file.

    Parameters
    ----------
    file_bytes : bytes
        Raw bytes of the DOCX file.

    Returns
    -------
    str
        Extracted text content.

    Raises
    ------
    ValueError
        If the DOCX cannot be read.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))

        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        extracted = "\n".join(paragraphs)

        if not extracted.strip():
            raise ValueError("No text could be extracted from this DOCX file.")

        logger.info(f"Extracted {len(extracted)} characters from DOCX ({len(doc.paragraphs)} paragraphs).")
        return extracted

    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}") from e


def validate_file(filename: str, content_type: str, file_bytes: bytes, max_size_mb: int = 10) -> None:
    """
    Validate uploaded file type and size.

    Parameters
    ----------
    filename : str
        Original filename (used to check extension).
    content_type : str
        MIME type reported by the client.
    file_bytes : bytes
        Raw file bytes (for size check).
    max_size_mb : int
        Maximum allowed file size in megabytes.

    Raises
    ------
    ValueError
        If the file fails any validation check.
    """
    allowed_extensions = {".pdf", ".docx"}
    allowed_mime_types = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    # Check extension
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in allowed_extensions:
        raise ValueError(f"Unsupported file type: '{ext}'. Only PDF and DOCX are accepted.")

    # Check MIME type
    if content_type not in allowed_mime_types:
        raise ValueError(f"Unsupported MIME type: '{content_type}'. Only PDF and DOCX are accepted.")

    # Check file size
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > max_size_mb:
        raise ValueError(f"File too large: {file_size_mb:.1f}MB. Maximum allowed is {max_size_mb}MB.")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Detect file type and dispatch to the correct extractor.

    Parameters
    ----------
    filename : str
        Original filename.
    file_bytes : bytes
        Raw file bytes.

    Returns
    -------
    str
        Extracted text content.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Cannot extract text from unsupported file type: '{ext}'.")
