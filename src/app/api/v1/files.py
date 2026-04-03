"""
File Download Endpoints

GET /api/v1/files/download/{report_id}/resume   — Download the AI-rewritten resume as PDF or DOCX
"""

from typing import Annotated, Literal

from fastapi import APIRouter, Depends
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession

from ...core.db.database import async_get_db
from ...core.exceptions.http_exceptions import ForbiddenException, NotFoundException
from ...crud.crud_report import crud_report
from ..dependencies import get_current_user

router = APIRouter(prefix="/files", tags=["File Downloads"])


@router.get("/download/{report_id}/resume")
async def download_optimized_resume(
    report_id: int,
    current_user: Annotated[dict, Depends(get_current_user)],
    db: Annotated[AsyncSession, Depends(async_get_db)],
    format: Literal["pdf", "docx"] = "pdf",
) -> Response:
    """
    Download the AI-rewritten (optimized) resume from a completed report.

    - **format**: Output format — `pdf` (default) or `docx`

    The optimized resume text is converted to the requested format on-the-fly.
    """
    report = await crud_report.get(db, id=report_id)
    if not report:
        raise NotFoundException("Report not found.")
    if report["user_id"] != current_user["id"]:
        raise ForbiddenException("You do not have access to this report.")

    optimized_text = report.get("optimized_resume_text")
    if not optimized_text:
        raise NotFoundException(
            "This report does not have an optimized resume yet. "
            "Please run /api/v1/rewrite/ first."
        )

    if format == "pdf":
        file_bytes = _generate_pdf(optimized_text)
        media_type = "application/pdf"
        filename = f"optimized_resume_{report_id}.pdf"
    else:
        file_bytes = _generate_docx(optimized_text)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"optimized_resume_{report_id}.docx"

    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _generate_pdf(text: str) -> bytes:
    """
    Generate a PDF from plain text using fpdf2.

    Raises
    ------
    RuntimeError
        If fpdf2 is not installed.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("fpdf2 is not installed. Run: pip install fpdf2")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.set_margins(left=20, top=20, right=20)

    for line in text.splitlines():
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue

        # Detect section headers (ALL CAPS lines or ending with ":")
        if line.isupper() or (len(line) < 50 and line.endswith(":")):
            pdf.set_font("Helvetica", style="B", size=12)
            pdf.multi_cell(0, 8, line)
            pdf.set_font("Helvetica", size=11)
        else:
            pdf.multi_cell(0, 6, line)

    return bytes(pdf.output())


def _generate_docx(text: str) -> bytes:
    """
    Generate a DOCX from plain text using python-docx.

    Raises
    ------
    RuntimeError
        If python-docx is not installed.
    """
    try:
        import io

        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in text.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Detect section headers
        if line.isupper() or (len(line) < 50 and line.endswith(":")):
            heading = doc.add_heading(line.rstrip(":"), level=2)
            heading.style.font.name = "Calibri"
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
